from docx import Document
from docx.shared import Inches
from datetime import datetime

full_name = input("enter you full name: ")

document = Document()
# document.add_picture('background.jpeg', width=Inches(2))
document.add_heading('Corporation Of Nothing', 2)
document.add_heading('Certificate of Absolute Shiet Programming SKills', 1)
document.add_paragraph('This is to certify that')
document.add_heading(full_name, 1)
document.add_paragraph("For demonstrating an exceptional ability to: Write code that runs nowhere, Debug problems that don’t exist (and miss the ones that do), Master the ancient art of Googling Stack Overflow for the wrong error, And commit code so meaningless, Git itself questioned its purpose.")
document.add_paragraph(f'Issued on {datetime.now().strftime("%d/%m/%Y")}')
document.add_paragraph('Ms Defjoy')
document.add_paragraph('CEO OF THE WORST PROGRAMMING SKILLS')
document.save('certificate.docx')